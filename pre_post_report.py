# -*- coding: utf-8 -*-
"""
Informe específico de respuesta aguda PRE-POST.

SportsLabResearch-BreastCancer-Wellbeing-Analyzer
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from src.analysis.moment_filter import pair_pre_post
from src.cleaning.clinical_cleaning import clean_clinical_data
from src.config.variable_mapping import prepare_form_dataframe
from src.core.session import get_current_participant


PROJECT_NAME = "SportsLabResearch-BreastCancer-Wellbeing-Analyzer"
ORGANIZATION = "SportsLabResearch"
AUTHOR = "José Pino-Ortega"

ROOT = Path(__file__).resolve().parents[2]
RESULTS_DIR = ROOT / "results"

NAVY = "12376B"
BLUE = "178FE5"
TEAL = "159A9C"
WHITE = "FFFFFF"
BLACK = "101828"
TEXT_GREY = "475467"
MID_GREY = "D0D5DD"
LIGHT_GREY = "F2F4F7"
LIGHT_BLUE = "EAF4FF"

VARIABLES: dict[str, tuple[str, str]] = {
    "sleep": ("Sueño", "puntos"),
    "mood": ("Estado de ánimo", "puntos"),
    "stress": ("Estrés", "puntos"),
    "fatigue": ("Fatiga", "puntos"),
    "upper_pain": ("Dolor superior", "puntos"),
    "lower_pain": ("Dolor inferior", "puntos"),
    "hr": ("Frecuencia cardiaca", "bpm"),
    "rmssd": ("RMSSD", "ms"),
    "ln_rmssd": ("LnRMSSD", "ln(ms)"),
    "sbp": ("Presión arterial sistólica", "mmHg"),
    "dbp": ("Presión arterial diastólica", "mmHg"),
    "spo2": ("Saturación de oxígeno", "%"),
}


# -----------------------------------------------------------------------------
# Utilidades generales
# -----------------------------------------------------------------------------


def safe_filename(value: object) -> str:
    text = str(value or "").strip()
    for old, new in {
        " ": "_", "/": "-", "\\": "-", ":": "-", "*": "",
        "?": "", '"': "", "<": "", ">": "", "|": "",
    }.items():
        text = text.replace(old, new)
    return text or "participante"


def format_value(value: Any, decimals: int = 2) -> str:
    if value is None or pd.isna(value):
        return "-"
    try:
        return f"{float(value):.{decimals}f}"
    except (TypeError, ValueError):
        return str(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = MID_GREY, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for name, value in (
        ("top", top), ("start", start), ("bottom", bottom), ("end", end)
    ):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    section.header_distance = Cm(0.45)
    section.footer_distance = Cm(0.45)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    for name, size, color in (
        ("Title", 24, NAVY),
        ("Heading 1", 15, NAVY),
        ("Heading 2", 12, BLUE),
    ):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(7)
        style.paragraph_format.space_after = Pt(4)


def add_header_footer(document: Document) -> None:
    for section in document.sections:
        header = section.header
        table = header.add_table(rows=1, cols=2, width=section.page_width)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        left, right = table.rows[0].cells
        for cell in (left, right):
            set_cell_border(cell, NAVY, "6")
            set_cell_margins(cell, top=15, bottom=65, start=20, end=20)

        p = left.paragraphs[0]
        run = p.add_run("SportsLab")
        run.bold = True
        run.font.name = "Aptos Display"
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor.from_string(NAVY)

        run = p.add_run("Research")
        run.bold = True
        run.font.name = "Aptos Display"
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor.from_string(BLUE)

        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("Respuesta aguda PRE–POST")
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(NAVY)

        footer = section.footer
        table = footer.add_table(rows=1, cols=3, width=section.page_width)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        for cell in table.rows[0].cells:
            set_cell_border(cell, NAVY, "4")
            set_cell_margins(cell, top=45, bottom=5, start=20, end=20)

        p = table.cell(0, 0).paragraphs[0]
        run = p.add_run(ORGANIZATION)
        run.bold = True
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor.from_string(NAVY)

        p = table.cell(0, 1).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(PROJECT_NAME)
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor.from_string(TEXT_GREY)

        p = table.cell(0, 2).paragraphs[0]
        add_page_number(p)
        for run in p.runs:
            run.font.size = Pt(7.5)
            run.font.color.rgb = RGBColor.from_string(TEXT_GREY)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


# -----------------------------------------------------------------------------
# Preparación del análisis
# -----------------------------------------------------------------------------


def effect_size_dz(pre: pd.Series, post: pd.Series) -> float:
    differences = pd.to_numeric(post, errors="coerce") - pd.to_numeric(
        pre, errors="coerce"
    )
    differences = differences.dropna()
    if len(differences) < 2:
        return np.nan
    standard_deviation = float(differences.std(ddof=1))
    if standard_deviation == 0:
        return np.nan
    return float(differences.mean() / standard_deviation)


def effect_label(value: float) -> str:
    if pd.isna(value):
        return "No calculable"
    magnitude = abs(float(value))
    if magnitude < 0.20:
        return "Trivial"
    if magnitude < 0.50:
        return "Pequeño"
    if magnitude < 0.80:
        return "Moderado"
    return "Grande"


def change_word(delta: float) -> str:
    if pd.isna(delta):
        return "sin estimación"
    if delta > 0:
        return "aumentó"
    if delta < 0:
        return "disminuyó"
    return "se mantuvo"


def build_analysis_table(clinical_data: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    available_variables = [
        variable for variable in VARIABLES if variable in clinical_data.columns
    ]
    if not available_variables:
        raise ValueError("No existen variables clínicas disponibles para comparar.")

    paired = pair_pre_post(clinical_data, available_variables)
    if paired.empty:
        raise ValueError(
            "No se encontraron sesiones con registros PRE y POST emparejables."
        )

    rows: list[dict[str, Any]] = []

    for variable in available_variables:
        pre_column = f"{variable}_pre"
        post_column = f"{variable}_post"
        if pre_column not in paired.columns or post_column not in paired.columns:
            continue

        valid = paired[[pre_column, post_column]].dropna()
        if valid.empty:
            continue

        pre_mean = float(valid[pre_column].mean())
        post_mean = float(valid[post_column].mean())
        delta = post_mean - pre_mean
        percent = np.nan if pre_mean == 0 else delta / abs(pre_mean) * 100
        dz = effect_size_dz(valid[pre_column], valid[post_column])
        label, unit = VARIABLES[variable]

        rows.append({
            "variable": variable,
            "label": label,
            "unit": unit,
            "n": int(len(valid)),
            "pre": pre_mean,
            "post": post_mean,
            "delta": delta,
            "percent": percent,
            "dz": dz,
            "effect": effect_label(dz),
        })

    analysis = pd.DataFrame(rows)
    if analysis.empty:
        raise ValueError("No existen pares PRE-POST válidos para las variables analizadas.")

    return paired, analysis


def session_dates(paired: pd.DataFrame) -> str:
    if "date" not in paired.columns:
        return "No disponible"
    dates = pd.to_datetime(paired["date"], errors="coerce", dayfirst=True).dropna()
    if dates.empty:
        return "No disponible"
    unique_dates = sorted(dates.dt.normalize().unique())
    if len(unique_dates) == 1:
        return pd.Timestamp(unique_dates[0]).strftime("%d/%m/%Y")
    return (
        pd.Timestamp(unique_dates[0]).strftime("%d/%m/%Y")
        + " – "
        + pd.Timestamp(unique_dates[-1]).strftime("%d/%m/%Y")
    )


# -----------------------------------------------------------------------------
# Contenido del informe
# -----------------------------------------------------------------------------


def add_cover(
    document: Document,
    participant: dict[str, Any],
    paired: pd.DataFrame,
) -> None:
    document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("INFORME DE RESPUESTA AGUDA")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(27)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("ANÁLISIS PRE–POST")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    details = [
        ("Identificador", participant.get("participant_id", "No disponible")),
        ("Participante", participant.get("name", "No disponible")),
        ("Sede", participant.get("site", "No disponible")),
        ("Sesiones emparejadas", len(paired)),
        ("Fechas analizadas", session_dates(paired)),
        ("Fecha de generación", datetime.now().strftime("%d/%m/%Y %H:%M")),
        ("Autor", AUTHOR),
    ]

    for index, (label, value) in enumerate(details):
        cells = table.add_row().cells
        set_cell_shading(cells[0], NAVY)
        set_cell_shading(cells[1], LIGHT_BLUE if index % 2 == 0 else WHITE)
        for cell in cells:
            set_cell_border(cell)
            set_cell_margins(cell, top=110, bottom=110, start=140, end=140)

        p = cells[0].paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)

        p = cells[1].paragraphs[0]
        p.add_run(str(value))

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Documento orientado a describir el cambio observado entre los registros "
        "anteriores y posteriores a la sesión."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(TEXT_GREY)

    document.add_page_break()


def add_visual_summary(
    document: Document,
    participant: dict[str, Any],
    paired: pd.DataFrame,
    analysis: pd.DataFrame,
) -> None:
    add_heading(document, "1. Respuesta aguda a la sesión (PRE–POST)")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"Participante: {participant.get('name', 'No disponible')}   ·   "
        f"Fechas: {session_dates(paired)}   ·   "
        f"Sesiones emparejadas: {len(paired)}"
    )
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["Variable", "PRE", "POST", "Cambio", "Interpretación"]

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, WHITE)
        set_cell_margins(cell, top=105, bottom=105)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)

    for row_index, item in analysis.iterrows():
        cells = table.add_row().cells
        fill = WHITE if row_index % 2 == 0 else LIGHT_GREY
        for cell in cells:
            set_cell_shading(cell, fill)
            set_cell_border(cell)
            set_cell_margins(cell, top=105, bottom=105, start=110, end=110)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        cells[0].text = str(item["label"])
        cells[1].text = f"●  {format_value(item['pre'])} {item['unit']}"
        cells[2].text = f"●  {format_value(item['post'])} {item['unit']}"
        cells[3].text = (
            f"{format_value(item['delta'])} {item['unit']}\n"
            f"({format_value(item['percent'])} %)"
        )
        cells[4].text = (
            f"{change_word(float(item['delta'])).capitalize()}\n"
            f"Magnitud: {item['effect']}"
        )

        for column_index, cell in enumerate(cells):
            p = cell.paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if column_index in {0, 4}
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            for run in p.runs:
                run.font.size = Pt(8.7)
                if column_index in {0, 1, 2}:
                    run.bold = True
                if column_index in {1, 2}:
                    run.font.color.rgb = RGBColor.from_string(BLUE)

        prevent_row_split(table.rows[-1])

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "La conexión PRE–POST representa el cambio medio de las sesiones emparejadas. "
        "No se asigna una valoración favorable o desfavorable de forma automática."
    )
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(TEXT_GREY)

    document.add_page_break()


def add_comparison_table(document: Document, analysis: pd.DataFrame) -> None:
    add_heading(document, "2. Comparación cuantitativa PRE–POST")

    table = document.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["Variable", "Unidad", "n", "PRE", "POST", "Δ", "Δ %", "dz"]

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, WHITE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)

    for row_index, item in analysis.iterrows():
        cells = table.add_row().cells
        values = [
            item["label"], item["unit"], str(item["n"]),
            format_value(item["pre"]), format_value(item["post"]),
            format_value(item["delta"]), format_value(item["percent"]),
            format_value(item["dz"]),
        ]
        fill = WHITE if row_index % 2 == 0 else LIGHT_GREY

        for column_index, value in enumerate(values):
            cell = cells[column_index]
            set_cell_shading(cell, fill)
            set_cell_border(cell)
            set_cell_margins(cell, top=90, bottom=90)
            p = cell.paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if column_index == 0
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            p.add_run(str(value))

        prevent_row_split(table.rows[-1])

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("Notas: ")
    run.bold = True
    p.add_run(
        "Δ = POST − PRE. Δ % = cambio relativo respecto al valor PRE. "
        "dz = tamaño del efecto estandarizado de las diferencias pareadas."
    )

    document.add_page_break()


def interpretation_text(item: pd.Series) -> str:
    direction = change_word(float(item["delta"]))
    if direction == "se mantuvo":
        first_sentence = (
            f"{item['label']} se mantuvo sin variación media entre PRE y POST."
        )
    else:
        first_sentence = (
            f"{item['label']} {direction} desde {format_value(item['pre'])} "
            f"hasta {format_value(item['post'])} {item['unit']}."
        )

    return (
        f"{first_sentence} El cambio medio fue de {format_value(item['delta'])} "
        f"{item['unit']} ({format_value(item['percent'])} %). "
        f"La magnitud estandarizada del cambio fue {str(item['effect']).lower()} "
        f"(dz = {format_value(item['dz'])}). Esta descripción no implica por sí sola "
        "una valoración favorable o desfavorable y debe interpretarse según el contexto "
        "de la sesión y las condiciones de medición."
    )


def add_interpretation(document: Document, analysis: pd.DataFrame) -> None:
    add_heading(document, "3. Interpretación de la respuesta")

    for row_index, item in analysis.iterrows():
        table = document.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        left, right = table.rows[0].cells

        set_cell_shading(left, NAVY)
        set_cell_border(left, NAVY)
        set_cell_margins(left, top=115, bottom=115, start=130, end=130)

        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(item["label"]))
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)

        p = left.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"PRE {format_value(item['pre'])}  →  POST {format_value(item['post'])}"
        )
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(WHITE)

        set_cell_shading(right, LIGHT_BLUE if row_index % 2 == 0 else WHITE)
        set_cell_border(right)
        set_cell_margins(right, top=115, bottom=115, start=140, end=140)
        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(interpretation_text(item))

        prevent_row_split(table.rows[0])
        document.add_paragraph().paragraph_format.space_after = Pt(1)

    document.add_page_break()


def group_summary(analysis: pd.DataFrame, variables: list[str]) -> str:
    selected = analysis[analysis["variable"].isin(variables)]
    if selected.empty:
        return "No existen datos emparejados suficientes para este bloque."

    fragments = []
    for _, item in selected.iterrows():
        fragments.append(
            f"{item['label']} {change_word(float(item['delta']))} "
            f"{format_value(abs(float(item['delta'])))} {item['unit']}"
        )
    return "; ".join(fragments) + "."


def add_conclusion(document: Document, analysis: pd.DataFrame) -> None:
    add_heading(document, "4. Conclusión de la sesión")

    blocks = [
        (
            "Respuesta cardiovascular",
            group_summary(analysis, ["hr", "sbp", "dbp", "spo2"]),
        ),
        (
            "Respuesta autonómica",
            group_summary(analysis, ["rmssd", "ln_rmssd"]),
        ),
        (
            "Bienestar percibido",
            group_summary(analysis, ["sleep", "mood", "stress", "fatigue"]),
        ),
        (
            "Dolor percibido",
            group_summary(analysis, ["upper_pain", "lower_pain"]),
        ),
    ]

    for title, text in blocks:
        table = document.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        left, right = table.rows[0].cells

        set_cell_shading(left, NAVY)
        set_cell_border(left, NAVY)
        set_cell_margins(left, top=130, bottom=130, start=140, end=140)
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)

        set_cell_shading(right, LIGHT_BLUE)
        set_cell_border(right)
        set_cell_margins(right, top=130, bottom=130, start=150, end=150)
        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(text)

        prevent_row_split(table.rows[0])
        document.add_paragraph().paragraph_format.space_after = Pt(2)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("Interpretación global: ")
    run.bold = True
    p.add_run(
        "el informe describe la respuesta media observada entre PRE y POST. "
        "La dirección de cada cambio no debe clasificarse automáticamente como favorable "
        "o desfavorable, ya que depende del objetivo de la sesión, la intensidad, el tiempo "
        "transcurrido hasta la medición POST, el estado clínico y la evolución individual."
    )

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("Consideración final: ")
    run.bold = True
    p.add_run(
        "este documento es una herramienta de apoyo al seguimiento y no sustituye la "
        "valoración clínica ni constituye un diagnóstico."
    )


# -----------------------------------------------------------------------------
# Generador público
# -----------------------------------------------------------------------------


def generate_pre_post_report(
    records: pd.DataFrame,
    output_dir: Path | str | None = None,
) -> Path:
    """Genera un único informe Word específico PRE-POST."""
    if records is None or records.empty:
        raise ValueError("No hay registros disponibles para generar el informe PRE-POST.")

    participant = get_current_participant() or {
        "participant_id": "sin_id",
        "name": "Participante",
        "site": "No disponible",
    }

    clinical_data = prepare_form_dataframe(records.copy())
    clinical_data = clean_clinical_data(clinical_data)

    paired, analysis = build_analysis_table(clinical_data)

    if output_dir is None:
        participant_folder = (
            RESULTS_DIR
            / safe_filename(participant.get("site"))
            / (
                safe_filename(participant.get("participant_id"))
                + "_"
                + safe_filename(participant.get("name"))
            )
        )
    else:
        participant_folder = Path(output_dir)

    participant_folder.mkdir(parents=True, exist_ok=True)

    output_path = participant_folder / (
        "Informe_PRE_POST_"
        + safe_filename(participant.get("participant_id"))
        + "_"
        + safe_filename(participant.get("name"))
        + ".docx"
    )

    document = Document()
    configure_document(document)
    add_header_footer(document)
    add_cover(document, participant, paired)
    add_visual_summary(document, participant, paired, analysis)
    add_comparison_table(document, analysis)
    add_interpretation(document, analysis)
    add_conclusion(document, analysis)

    document.core_properties.title = "Informe de respuesta aguda PRE-POST"
    document.core_properties.subject = PROJECT_NAME
    document.core_properties.author = AUTHOR
    document.core_properties.keywords = (
        "cáncer de mama, bienestar, respuesta aguda, PRE-POST, ejercicio"
    )
    document.save(output_path)

    return output_path


if __name__ == "__main__":
    raise SystemExit(
        "Este módulo debe ejecutarse desde main.py con una participante activa."
    )
