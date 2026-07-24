# -*- coding: utf-8 -*-

"""
SportsLabResearch-BreastCancer-Wellbeing-Analyzer

Bloques Word para el análisis por momento de registro.

Genera:

    - Estado PRE
    - Estado POST
    - Efecto PRE-POST
    - Interpretación clínica automática

Este módulo no crea el documento Word completo.
Añade las secciones a un objeto Document ya existente.
"""

from __future__ import annotations

from typing import Any

import pandas as pd

from src.analysis.moment_analysis import run_moment_analysis
from src.analysis.pre_post import (
    MODE_POST,
    MODE_PRE,
    MODE_PRE_POST,
)


def safe_number(
    value: Any,
    decimals: int = 2,
) -> str:
    """
    Formatea un valor numérico de forma segura.
    """

    try:
        if pd.isna(value):
            return "-"
    except (TypeError, ValueError):
        pass

    try:
        return f"{float(value):.{decimals}f}"
    except (TypeError, ValueError):
        return str(value)


def set_cell_text(
    cell: Any,
    value: Any,
) -> None:
    """
    Escribe texto dentro de una celda Word.
    """

    cell.text = str(value)


def add_table_header(
    table: Any,
    headers: list[str],
) -> None:
    """
    Añade los encabezados de una tabla.
    """

    header_cells = table.rows[0].cells

    for index, header in enumerate(headers):
        set_cell_text(
            header_cells[index],
            header,
        )


def add_single_moment_table(
    document: Any,
    results: pd.DataFrame,
    title: str,
    description: str,
) -> None:
    """
    Añade una tabla descriptiva PRE o POST.
    """

    document.add_heading(
        title,
        level=2,
    )

    document.add_paragraph(
        description
    )

    if results is None or results.empty:
        document.add_paragraph(
            "No hay datos disponibles para este momento de registro."
        )
        return

    headers = [
        "Variable",
        "n",
        "Media",
        "Mediana",
        "DE",
        "Mínimo",
        "Máximo",
    ]

    table = document.add_table(
        rows=1,
        cols=len(headers),
    )

    table.style = "Table Grid"

    add_table_header(
        table,
        headers,
    )

    for _, row in results.iterrows():

        cells = table.add_row().cells

        values = [
            row.get(
                "variable_label",
                row.get("variable", "-"),
            ),
            row.get("n", "-"),
            safe_number(row.get("mean")),
            safe_number(row.get("median")),
            safe_number(row.get("sd")),
            safe_number(row.get("minimum")),
            safe_number(row.get("maximum")),
        ]

        for index, value in enumerate(values):
            set_cell_text(
                cells[index],
                value,
            )


def variation_direction_text(
    delta: Any,
) -> str:
    """
    Describe la dirección matemática de la variación.
    """

    try:
        numeric_delta = float(delta)
    except (TypeError, ValueError):
        return "no pudo determinarse"

    if pd.isna(numeric_delta):
        return "no pudo determinarse"

    if numeric_delta > 0:
        return "aumentó"

    if numeric_delta < 0:
        return "disminuyó"

    return "no cambió"


def clinical_interpretation_text(
    row: pd.Series,
) -> str:
    """
    Genera una interpretación clínica breve.
    """

    variable = row.get(
        "variable_label",
        row.get("variable", "La variable"),
    )

    pre_value = safe_number(
        row.get("before_mean")
    )

    post_value = safe_number(
        row.get("after_mean")
    )

    delta = row.get("delta")
    delta_text = safe_number(delta)

    delta_percent = safe_number(
        row.get("delta_percent")
    )

    direction = variation_direction_text(
        delta
    )

    interpretation = row.get(
        "change_interpretation",
        "No evaluable",
    )

    magnitude = row.get(
        "effect_magnitude",
        "No evaluable",
    )

    if interpretation == "Cambio favorable":
        clinical_sentence = (
            "La dirección del cambio se interpreta como favorable."
        )

    elif interpretation == "Cambio desfavorable":
        clinical_sentence = (
            "La dirección del cambio se interpreta como desfavorable."
        )

    elif interpretation == "Sin cambio relevante":
        clinical_sentence = (
            "No se observa un cambio clínicamente relevante."
        )

    else:
        clinical_sentence = (
            "La relevancia clínica del cambio no pudo evaluarse."
        )

    if magnitude == "No evaluable":
        magnitude_sentence = (
            "La magnitud del efecto no pudo calcularse."
        )
    else:
        magnitude_sentence = (
            f"La magnitud del efecto fue {str(magnitude).lower()}."
        )

    return (
        f"{variable} {direction} desde {pre_value} en PRE "
        f"hasta {post_value} en POST. "
        f"La variación media fue de {delta_text} puntos "
        f"({delta_percent} %). "
        f"{clinical_sentence} "
        f"{magnitude_sentence}"
    )


def add_pre_post_table(
    document: Any,
    results: pd.DataFrame,
) -> None:
    """
    Añade la tabla comparativa PRE-POST.
    """

    document.add_heading(
        "Efecto PRE–POST",
        level=2,
    )

    document.add_paragraph(
        "Este análisis compara los registros realizados antes "
        "y después de la intervención. La variación se calcula "
        "como POST menos PRE."
    )

    if results is None or results.empty:
        document.add_paragraph(
            "No existen pares PRE–POST válidos para las "
            "variables disponibles."
        )
        return

    headers = [
        "Variable",
        "n",
        "PRE",
        "POST",
        "Variación",
        "Variación %",
        "p",
        "dz",
        "Magnitud",
        "Interpretación",
    ]

    table = document.add_table(
        rows=1,
        cols=len(headers),
    )

    table.style = "Table Grid"

    add_table_header(
        table,
        headers,
    )

    for _, row in results.iterrows():

        cells = table.add_row().cells

        values = [
            row.get(
                "variable_label",
                row.get("variable", "-"),
            ),
            row.get("n", "-"),
            safe_number(row.get("before_mean")),
            safe_number(row.get("after_mean")),
            safe_number(row.get("delta")),
            safe_number(row.get("delta_percent")),
            safe_number(row.get("p"), decimals=3),
            safe_number(row.get("cohen_dz")),
            row.get("effect_magnitude", "-"),
            row.get("change_interpretation", "-"),
        ]

        for index, value in enumerate(values):
            set_cell_text(
                cells[index],
                value,
            )


def add_clinical_interpretation(
    document: Any,
    results: pd.DataFrame,
) -> None:
    """
    Añade la interpretación clínica narrativa.
    """

    document.add_heading(
        "Interpretación clínica del efecto",
        level=2,
    )

    if results is None or results.empty:
        document.add_paragraph(
            "No fue posible generar una interpretación PRE–POST "
            "por ausencia de pares válidos."
        )
        return

    for _, row in results.iterrows():

        paragraph = document.add_paragraph(
            style="List Bullet",
        )

        paragraph.add_run(
            clinical_interpretation_text(row)
        )


def add_moment_report_sections(
    document: Any,
    dataframe: pd.DataFrame,
) -> dict[str, pd.DataFrame]:
    """
    Añade al documento las secciones PRE, POST y PRE-POST.

    Parameters
    ----------
    document:
        Documento Word de python-docx.

    dataframe:
        Datos clínicos preparados, con la columna moment.

    Returns
    -------
    dict:
        Resultados PRE, POST y PRE_POST.
    """

    document.add_heading(
        "Análisis por momento de registro",
        level=1,
    )

    document.add_paragraph(
        "Los registros se clasifican según el momento en que "
        "fueron cumplimentados: 1 corresponde al registro PRE "
        "y 2 al registro POST. Además, se analiza el efecto "
        "agudo mediante la comparación pareada PRE–POST."
    )

    if dataframe is None or dataframe.empty:

        document.add_paragraph(
            "No hay registros disponibles para realizar "
            "el análisis por momento."
        )

        return {
            MODE_PRE: pd.DataFrame(),
            MODE_POST: pd.DataFrame(),
            MODE_PRE_POST: pd.DataFrame(),
        }

    pre_results = run_moment_analysis(
        dataframe=dataframe,
        mode=MODE_PRE,
    )

    post_results = run_moment_analysis(
        dataframe=dataframe,
        mode=MODE_POST,
    )

    pre_post_results = run_moment_analysis(
        dataframe=dataframe,
        mode=MODE_PRE_POST,
    )

    add_single_moment_table(
        document=document,
        results=pre_results,
        title="Estado PRE",
        description=(
            "El estado PRE representa la situación de la "
            "participante antes de la intervención o sesión."
        ),
    )

    add_single_moment_table(
        document=document,
        results=post_results,
        title="Estado POST",
        description=(
            "El estado POST representa la situación de la "
            "participante después de la intervención o sesión."
        ),
    )

    add_pre_post_table(
        document=document,
        results=pre_post_results,
    )

    add_clinical_interpretation(
        document=document,
        results=pre_post_results,
    )

    return {
        MODE_PRE: pre_results,
        MODE_POST: post_results,
        MODE_PRE_POST: pre_post_results,
    }


if __name__ == "__main__":

    from docx import Document

    sample = pd.DataFrame({
        "participant": [
            "A", "A",
            "B", "B",
            "C", "C",
        ],
        "date": pd.to_datetime([
            "2026-07-01",
            "2026-07-01",
            "2026-07-02",
            "2026-07-02",
            "2026-07-03",
            "2026-07-03",
        ]),
        "moment": [
            1, 2,
            1, 2,
            1, 2,
        ],
        "fatigue": [
            6, 4,
            5, 4,
            7, 5,
        ],
        "energy": [
            4, 6,
            5, 6,
            3, 5,
        ],
        "stress": [
            7, 5,
            6, 4,
            5, 3,
        ],
    })

    output = Document()

    output.add_heading(
        "Prueba del análisis por momento",
        level=0,
    )

    add_moment_report_sections(
        document=output,
        dataframe=sample,
    )

    output.save(
        "results/moment_report_test.docx"
    )

    print(
        "Informe de prueba generado: "
        "results/moment_report_test.docx"
    )
