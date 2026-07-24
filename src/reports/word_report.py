# -*- coding: utf-8 -*-
"""
SportsLabResearch-BreastCancer-Wellbeing-Analyzer
Generador automático de informes clínicos en Word.

Ejecución directa desde PowerShell:
    py .\\src\\reports\\word_report.py

Dependencia:
    py -m pip install python-docx
"""

from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path
from typing import Any, Mapping

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PROJECT_NAME = "SportsLabResearch-BreastCancer-Wellbeing-Analyzer"
NAVY = "12376B"
BLUE = "178FE5"
LIGHT_BLUE = "DCEEFF"
LIGHT_GREY = "F2F4F7"
MID_GREY = "D0D5DD"
TEXT_GREY = "475467"
GREEN = "2E8B43"
LIGHT_GREEN = "EAF6EC"
WHITE = "FFFFFF"
BLACK = "111111"


def project_root() -> Path:
    current = Path(__file__).resolve()
    for parent in [current.parent, *current.parents]:
        if (parent / "src").exists():
            return parent
    return Path.cwd().resolve()


def safe_text(value: Any, default: str = "No disponible") -> str:
    if value is None:
        return default
    text = str(value).strip()
    return text if text else default


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = MID_GREY, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(BLACK)


def add_header(document: Document) -> None:
    section = document.sections[0]
    table = section.header.add_table(rows=1, cols=2, width=section.page_width)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = table.rows[0].cells
    for cell in (left, right):
        set_cell_border(cell, NAVY, "6")
        set_cell_margins(cell, top=20, bottom=70, start=30, end=30)
    p = left.paragraphs[0]
    run = p.add_run("SportsLab")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    run = p.add_run("Research")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    p2 = right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p2.add_run("Breast Cancer Wellbeing Analyzer")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(NAVY)


def add_footer(document: Document) -> None:
    section = document.sections[0]
    table = section.footer.add_table(rows=1, cols=2, width=section.page_width)
    table.autofit = False
    left, right = table.rows[0].cells
    for cell in (left, right):
        set_cell_border(cell, NAVY, "4")
        set_cell_margins(cell, top=50, bottom=10)
    p = left.paragraphs[0]
    run = p.add_run(f"SportsLabResearch | {PROJECT_NAME}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(TEXT_GREY)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(datetime.now().strftime("Generado el %d/%m/%Y"))
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(TEXT_GREY)


def add_title(document: Document, participant: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INFORME CLÍNICO DE SEGUIMIENTO")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Participante: {safe_text(participant)}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(TEXT_GREY)


def add_section_title(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(NAVY)


def add_summary_cards(document: Document, summary: Mapping[str, Any]) -> None:
    cards = [
        ("Media", summary.get("media", "82 bpm")),
        ("Último valor", summary.get("ultimo_valor", "78 bpm")),
        ("Variación", summary.get("variacion", "-5 %")),
        ("Registros", summary.get("registros", "8")),
        ("Desviación estándar", summary.get("desviacion_estandar", "4.2 bpm")),
        ("Estado clínico", summary.get("estado_clinico", "ESTABLE")),
    ]
    table = document.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, (label, value) in enumerate(cards):
        row = index // 3
        col = index % 3
        cell = table.cell(row, col)
        fill = LIGHT_GREEN if label == "Estado clínico" else LIGHT_BLUE
        accent = GREEN if label == "Estado clínico" else NAVY
        set_cell_shading(cell, fill)
        set_cell_border(cell, accent, "6")
        set_cell_margins(cell, top=120, bottom=120, start=140, end=140)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(TEXT_GREY)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(safe_text(value))
        run.bold = True
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor.from_string(accent)


def add_chart(document: Document, chart_path: Path | None) -> None:
    add_section_title(document, "Evolución longitudinal")
    if chart_path and chart_path.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(chart_path), width=Cm(24.5))
    else:
        table = document.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        set_cell_shading(cell, LIGHT_GREY)
        set_cell_border(cell)
        set_cell_margins(cell, top=300, bottom=300)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("No se ha encontrado el gráfico clínico.")
        run.italic = True
        run.font.color.rgb = RGBColor.from_string(TEXT_GREY)


def add_interpretation(document: Document) -> None:
    add_section_title(document, "Interpretación clínica")
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GREEN)
    set_cell_border(cell, GREEN, "6")
    set_cell_margins(cell, top=150, bottom=150, start=170, end=170)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "La frecuencia cardiaca presenta una evolución estable durante el periodo analizado. "
        "El último registro se mantiene dentro del rango clínico normal y no se observa una "
        "tendencia sostenida al incremento. Se recomienda continuar el seguimiento longitudinal "
        "en condiciones de medición comparables."
    )


def add_ranges_table(document: Document) -> None:
    add_section_title(document, "Rangos clínicos de referencia")
    rows = [
        ("Normal", "50–89 bpm", "Rango esperado en reposo."),
        ("Alerta", "90–100 bpm", "Valores elevados; monitorizar tendencia."),
        ("Riesgo", ">100 bpm", "Valorar intervención clínica."),
    ]
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ("Clasificación", "Rango", "Interpretación")
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, WHITE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        for col_index, value in enumerate(values):
            cell = row.cells[col_index]
            set_cell_shading(cell, WHITE if row_index % 2 == 0 else LIGHT_GREY)
            set_cell_border(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(value)


def build_word_report(
    output_path: str | Path,
    participant: str = "Participante",
    chart_path: str | Path | None = None,
    summary: Mapping[str, Any] | None = None,
) -> Path:
    output = Path(output_path).resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    chart = Path(chart_path).resolve() if chart_path else None
    document = Document()
    configure_document(document)
    add_header(document)
    add_footer(document)
    add_title(document, participant)
    add_section_title(document, "Resumen de resultados")
    add_summary_cards(document, summary or {})
    add_chart(document, chart)
    add_interpretation(document)
    add_ranges_table(document)
    document.save(output)
    return output


def find_chart(root: Path) -> Path | None:
    candidates = [
        root / "results" / "charts" / "frecuencia_cardiaca.png",
        root / "results" / "charts" / "heart_rate.png",
        root / "results" / "frecuencia_cardiaca.png",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def main() -> None:
    root = project_root()
    parser = argparse.ArgumentParser()
    parser.add_argument("--participant", default="Participante")
    parser.add_argument("--chart", default=None)
    parser.add_argument(
        "--output",
        default=str(root / "results" / "reports" / "Informe_clinico.docx"),
    )
    args = parser.parse_args()
    chart = Path(args.chart) if args.chart else find_chart(root)
    output = build_word_report(
        output_path=args.output,
        participant=args.participant,
        chart_path=chart,
    )
    print(f"Informe Word generado correctamente: {output}")


if __name__ == "__main__":
    main()
