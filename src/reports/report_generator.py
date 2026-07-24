# -*- coding: utf-8 -*-
"""
Generación automática de informes Word individuales.

SportsLabResearch-BreastCancer-Wellbeing-Analyzer
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from src.analysis.blood_pressure import analyse as analyse_blood_pressure
from src.analysis.wellbeing import analyse as analyse_wellbeing
from src.cleaning.clinical_cleaning import clean_clinical_data
from src.config.variable_mapping import prepare_form_dataframe
from src.core.session import get_current_participant
from src.graphics.clinical_graphs import generate_clinical_graphs
from src.reports.moment_report import add_moment_report_sections


PROJECT_NAME = "SportsLabResearch-BreastCancer-Wellbeing-Analyzer"
ORGANIZATION = "SportsLabResearch"
AUTHOR = "José Pino-Ortega"

ROOT = Path(__file__).resolve().parents[2]
RESULTS_DIR = ROOT / "results"

REPORT_MODE_PRE = "PRE"
REPORT_MODE_POST = "POST"
REPORT_MODE_PRE_POST = "PRE_POST"
REPORT_MODES = {
    REPORT_MODE_PRE,
    REPORT_MODE_POST,
    REPORT_MODE_PRE_POST,
}

NAVY = "12376B"
BLUE = "178FE5"
TEAL = "159A9C"
GREEN = "2E8B43"
AMBER = "B77900"
RED = "D92D20"
WHITE = "FFFFFF"
BLACK = "101828"
TEXT_GREY = "475467"
MID_GREY = "D0D5DD"
LIGHT_GREY = "F2F4F7"
LIGHT_BLUE = "EAF4FF"
LIGHT_GREEN = "EAF6EC"
LIGHT_AMBER = "FFF4D6"
LIGHT_RED = "FDE8E7"

VARIABLES = {
    "hr": ("Frecuencia cardiaca", "bpm"),
    "rmssd": ("RMSSD", "ms"),
    "ln_rmssd": ("LnRMSSD", "ln(ms)"),
    "sbp": ("Presión arterial sistólica", "mmHg"),
    "dbp": ("Presión arterial diastólica", "mmHg"),
    "spo2": ("Saturación de oxígeno", "%"),
    "sleep": ("Sueño", "puntos"),
    "mood": ("Estado de ánimo", "puntos"),
    "stress": ("Estrés", "puntos"),
    "fatigue": ("Fatiga", "puntos"),
    "upper_pain": ("Dolor superior", "puntos"),
    "lower_pain": ("Dolor inferior", "puntos"),
}

GRAPH_TITLES = {
    "presion_arterial.png": "Evolución de la presión arterial",
    "frecuencia_cardiaca.png": "Evolución de la frecuencia cardiaca",
    "rmssd.png": "Evolución de RMSSD",
    "ln_rmssd.png": "Evolución de LnRMSSD",
    "spo2.png": "Evolución de la saturación de oxígeno",
    "sueno.png": "Evolución del sueño",
    "estado_animo.png": "Evolución del estado de ánimo",
    "estres.png": "Evolución del estrés",
    "fatiga.png": "Evolución de la fatiga",
    "dolor_superior.png": "Evolución del dolor superior",
    "dolor_inferior.png": "Evolución del dolor inferior",
}

GRAPH_VARIABLES = {
    "presion_arterial.png": ("sbp", "dbp"),
    "frecuencia_cardiaca.png": ("hr",),
    "rmssd.png": ("rmssd",),
    "ln_rmssd.png": ("ln_rmssd",),
    "spo2.png": ("spo2",),
    "sueno.png": ("sleep",),
    "estado_animo.png": ("mood",),
    "estres.png": ("stress",),
    "fatiga.png": ("fatigue",),
    "dolor_superior.png": ("upper_pain",),
    "dolor_inferior.png": ("lower_pain",),
}


def safe_filename(value: object) -> str:
    text = str(value or "").strip()
    for old, new in {
        " ": "_", "/": "-", "\\": "-", ":": "-", "*": "",
        "?": "", '"': "", "<": "", ">": "", "|": "",
    }.items():
        text = text.replace(old, new)
    return text or "participante"


def format_value(value: Any, decimals: int = 2) -> str:
    if value is None or pd.isna(value):
        return "-"
    try:
        return f"{float(value):.{decimals}f}"
    except (TypeError, ValueError):
        return str(value)


def numeric_values(dataframe: pd.DataFrame, column: str) -> pd.Series:
    if column not in dataframe.columns:
        return pd.Series(dtype="float64")
    return pd.to_numeric(dataframe[column], errors="coerce").dropna()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = MID_GREY, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)
    section.header_distance = Cm(0.45)
    section.footer_distance = Cm(0.45)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    for name, size, color in (
        ("Title", 24, NAVY),
        ("Heading 1", 15, NAVY),
        ("Heading 2", 12, BLUE),
        ("Heading 3", 10.5, TEAL),
    ):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(7)
        style.paragraph_format.space_after = Pt(4)


def add_header_footer(document: Document) -> None:
    for section in document.sections:
        header = section.header
        table = header.add_table(rows=1, cols=2, width=section.page_width)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        left, right = table.rows[0].cells
        for cell in (left, right):
            set_cell_border(cell, NAVY, "6")
            set_cell_margins(cell, top=15, bottom=65, start=20, end=20)

        p = left.paragraphs[0]
        run = p.add_run("SportsLab")
        run.bold = True
        run.font.name = "Aptos Display"
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor.from_string(NAVY)
        run = p.add_run("Research")
        run.bold = True
        run.font.name = "Aptos Display"
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor.from_string(BLUE)

        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("Breast Cancer Wellbeing Analyzer")
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(NAVY)

        footer = section.footer
        table = footer.add_table(rows=1, cols=3, width=section.page_width)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell in table.rows[0].cells:
            set_cell_border(cell, NAVY, "4")
            set_cell_margins(cell, top=45, bottom=5, start=20, end=20)

        p = table.cell(0, 0).paragraphs[0]
        run = p.add_run(ORGANIZATION)
        run.bold = True
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor.from_string(NAVY)

        p = table.cell(0, 1).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(PROJECT_NAME)
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor.from_string(TEXT_GREY)

        p = table.cell(0, 2).paragraphs[0]
        add_page_number(p)
        for run in p.runs:
            run.font.size = Pt(7.5)
            run.font.color.rgb = RGBColor.from_string(TEXT_GREY)


def add_section_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_cover(document: Document, participant: dict[str, Any], records: pd.DataFrame) -> None:
    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("INFORME INDIVIDUAL")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("BIENESTAR, SALUD Y SEGUIMIENTO LONGITUDINAL")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    details = [
        ("Identificador", participant.get("participant_id", "No disponible")),
        ("Participante", participant.get("name", "No disponible")),
        ("Sede", participant.get("site", "No disponible")),
        ("Registros incluidos", len(records)),
        ("Fecha de generación", datetime.now().strftime("%d/%m/%Y %H:%M")),
        ("Autor", AUTHOR),
    ]
    for index, (label, value) in enumerate(details):
        cells = table.add_row().cells
        set_cell_shading(cells[0], NAVY)
        set_cell_shading(cells[1], LIGHT_BLUE if index % 2 == 0 else WHITE)
        for cell in cells:
            set_cell_border(cell)
            set_cell_margins(cell, top=110, bottom=110, start=140, end=140)
        p = cells[0].paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        p = cells[1].paragraphs[0]
        run = p.add_run(str(value))
        run.font.color.rgb = RGBColor.from_string(BLACK)

    document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Documento de apoyo al seguimiento. Los resultados deben interpretarse junto con "
        "el contexto clínico y la valoración del equipo profesional."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(TEXT_GREY)
    document.add_page_break()


def dates_summary(clinical_data: pd.DataFrame) -> tuple[str, str]:
    if "date" not in clinical_data.columns:
        return "No disponible", "No disponible"
    dates = pd.to_datetime(clinical_data["date"], errors="coerce", dayfirst=True)
    if not dates.notna().any():
        return "No disponible", "No disponible"
    return dates.min().strftime("%d/%m/%Y"), dates.max().strftime("%d/%m/%Y")


def variable_status(column: str, last: float) -> tuple[str, str, str]:
    if column == "hr":
        if last > 100:
            return "RIESGO", RED, LIGHT_RED
        if last >= 90:
            return "SEGUIMIENTO", AMBER, LIGHT_AMBER
        return "FAVORABLE", GREEN, LIGHT_GREEN
    if column == "spo2":
        if last < 92:
            return "RIESGO", RED, LIGHT_RED
        if last < 95:
            return "SEGUIMIENTO", AMBER, LIGHT_AMBER
        return "FAVORABLE", GREEN, LIGHT_GREEN
    if column == "sleep":
        if last < 4:
            return "DESFAVORABLE", RED, LIGHT_RED
        if last < 7:
            return "SEGUIMIENTO", AMBER, LIGHT_AMBER
        return "FAVORABLE", GREEN, LIGHT_GREEN
    if column == "mood":
        if last <= 2:
            return "DESFAVORABLE", RED, LIGHT_RED
        if last == 3:
            return "SEGUIMIENTO", AMBER, LIGHT_AMBER
        return "FAVORABLE", GREEN, LIGHT_GREEN
    if column in {"stress", "fatigue", "upper_pain", "lower_pain"}:
        if last >= 7:
            return "DESFAVORABLE", RED, LIGHT_RED
        if last >= 4:
            return "SEGUIMIENTO", AMBER, LIGHT_AMBER
        return "FAVORABLE", GREEN, LIGHT_GREEN
    return "INFORMATIVO", NAVY, LIGHT_BLUE


def build_executive_items(clinical_data: pd.DataFrame) -> list[dict[str, str]]:
    priority = ["hr", "rmssd", "sbp", "dbp", "spo2", "sleep", "stress", "fatigue", "upper_pain", "lower_pain"]
    items: list[dict[str, str]] = []
    for column in priority:
        values = numeric_values(clinical_data, column)
        if values.empty:
            continue
        label, unit = VARIABLES[column]
        last = float(values.iloc[-1])
        mean = float(values.mean())
        status, color, fill = variable_status(column, last)
        items.append({
            "label": label,
            "value": f"{last:.1f} {unit}",
            "detail": f"Media {mean:.1f} · n={len(values)}",
            "status": status,
            "color": color,
            "fill": fill,
        })
    return items[:6]


def add_executive_summary(document: Document, clinical_data: pd.DataFrame) -> None:
    add_section_heading(document, "1. Resumen clínico ejecutivo")
    items = build_executive_items(clinical_data)
    if not items:
        document.add_paragraph("No existen datos suficientes para construir el resumen ejecutivo.")
        return

    rows = (len(items) + 2) // 3
    table = document.add_table(rows=rows, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, item in enumerate(items):
        cell = table.cell(index // 3, index % 3)
        set_cell_shading(cell, item["fill"])
        set_cell_border(cell, item["color"], "8")
        set_cell_margins(cell, top=115, bottom=115, start=130, end=130)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item["label"])
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(TEXT_GREY)

        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(item["value"])
        run.bold = True
        run.font.name = "Aptos Display"
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor.from_string(item["color"])

        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item["detail"])
        run.font.size = Pt(7.8)
        run.font.color.rgb = RGBColor.from_string(TEXT_GREY)

        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item["status"])
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(item["color"])


def add_introduction_and_participant(document: Document, participant: dict[str, Any], clinical_data: pd.DataFrame) -> None:
    add_section_heading(document, "2. Contexto del informe")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "Este informe sintetiza los registros individuales disponibles para describir el estado actual, "
        "la evolución temporal y la disponibilidad de datos de la participante. La lectura prioriza las "
        "tendencias longitudinales y los cambios persistentes frente a valores aislados."
    )

    first_date, last_date = dates_summary(clinical_data)
    table = document.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    details = [
        ("Identificador", participant.get("participant_id", "No disponible")),
        ("Sede", participant.get("site", "No disponible")),
        ("Primer registro", first_date),
        ("Último registro", last_date),
        ("Registros", len(clinical_data)),
        ("Variables válidas", sum(not numeric_values(clinical_data, c).empty for c in VARIABLES)),
        ("Generado", datetime.now().strftime("%d/%m/%Y")),
        ("Responsable", AUTHOR),
    ]
    for index in range(0, len(details), 4):
        cells = table.add_row().cells
        for col, (label, value) in enumerate(details[index:index + 4]):
            cell = cells[col]
            set_cell_shading(cell, LIGHT_GREY)
            set_cell_border(cell)
            set_cell_margins(cell, top=100, bottom=100, start=100, end=100)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(label + "\n")
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(NAVY)
            run = p.add_run(str(value))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(BLACK)


def style_table(table, header_fill: str = NAVY) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if table.rows:
        set_repeat_table_header(table.rows[0])
        for cell in table.rows[0].cells:
            set_cell_shading(cell, header_fill)
            set_cell_border(cell, WHITE)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8.3)
                run.font.color.rgb = RGBColor.from_string(WHITE)
    for row_index, row in enumerate(table.rows[1:], start=1):
        prevent_row_split(row)
        fill = WHITE if row_index % 2 else LIGHT_GREY
        for cell in row.cells:
            set_cell_shading(cell, fill)
            set_cell_border(cell)
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.2)


def add_descriptive_summary(document: Document, clinical_data: pd.DataFrame) -> None:
    add_section_heading(document, "3. Resumen descriptivo")
    table = document.add_table(rows=1, cols=8)
    headers = ["Variable", "Unidad", "n", "Media", "DE", "Mínimo", "Máximo", "Último"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    available = 0
    for column, (label, unit) in VARIABLES.items():
        values = numeric_values(clinical_data, column)
        if values.empty:
            continue
        available += 1
        cells = table.add_row().cells
        values_out = [
            label, unit, str(len(values)), format_value(values.mean()),
            format_value(values.std()), format_value(values.min()),
            format_value(values.max()), format_value(values.iloc[-1]),
        ]
        for i, value in enumerate(values_out):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER

    if available == 0:
        document.add_paragraph("No se encontraron variables clínicas válidas para el resumen.")
        return
    style_table(table)


def add_blood_pressure_section(document: Document, clinical_data: pd.DataFrame) -> None:
    add_section_heading(document, "4. Presión arterial")
    if numeric_values(clinical_data, "sbp").empty and numeric_values(clinical_data, "dbp").empty:
        document.add_paragraph("No hay datos de presión arterial disponibles.")
        return
    try:
        results = analyse_blood_pressure(clinical_data)
        summary = results.get("summary", {})
    except Exception as exc:
        document.add_paragraph(f"No se pudo completar el análisis de presión arterial: {exc}")
        return

    table = document.add_table(rows=2, cols=4)
    labels = ["PAS media", "PAD media", "Clasificación", "Registros válidos"]
    values = [
        f"{format_value(summary.get('sbp_mean'))} mmHg",
        f"{format_value(summary.get('dbp_mean'))} mmHg",
        str(summary.get("classification", "No disponible")),
        f"PAS {summary.get('valid_sbp', 0)} · PAD {summary.get('valid_dbp', 0)}",
    ]
    for i, label in enumerate(labels):
        cell = table.cell(0, i)
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, WHITE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        cell = table.cell(1, i)
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(values[i])
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(NAVY)

    interpretation = summary.get("interpretation", "No disponible")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("Interpretación: ")
    run.bold = True
    p.add_run(str(interpretation))


def add_wellbeing_section(document: Document, clinical_data: pd.DataFrame) -> None:
    add_section_heading(document, "5. Bienestar multidimensional")
    try:
        results = analyse_wellbeing(clinical_data)
    except Exception as exc:
        document.add_paragraph(f"No se pudo completar el análisis de bienestar: {exc}")
        return
    if results is None or results.empty:
        document.add_paragraph("No hay datos de bienestar disponibles.")
        return

    table = document.add_table(rows=1, cols=len(results.columns))
    for i, column in enumerate(results.columns):
        table.rows[0].cells[i].text = str(column)
    for _, row_data in results.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = format_value(value)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_table(table, TEAL)

    alerts: list[str] = []
    for column in ("sleep", "mood", "stress", "fatigue", "upper_pain", "lower_pain"):
        values = numeric_values(clinical_data, column)
        if values.empty:
            continue
        status, _, _ = variable_status(column, float(values.iloc[-1]))
        if status in {"DESFAVORABLE", "RIESGO"}:
            alerts.append(f"{VARIABLES[column][0]}: último valor {values.iloc[-1]:.1f}.")
    if alerts:
        table = document.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        set_cell_shading(cell, LIGHT_RED)
        set_cell_border(cell, RED, "7")
        set_cell_margins(cell, top=120, bottom=120, start=150, end=150)
        p = cell.paragraphs[0]
        run = p.add_run("Aspectos que requieren atención\n")
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(RED)
        p.add_run(" ".join(alerts))


def add_scientific_context(document: Document, clinical_data: pd.DataFrame) -> None:
    try:
        from src.evidence.evidence_engine import ScientificEvidenceEngine
    except Exception:
        return

    mapping = {
        "hr": "frecuencia cardíaca", "rmssd": "rmssd", "ln_rmssd": "ln_rmssd",
        "sbp": "presión arterial sistólica", "dbp": "presión arterial diastólica",
        "spo2": "saturación de oxígeno", "sleep": "sueño", "mood": "estado de ánimo",
        "stress": "estrés", "fatigue": "fatiga", "upper_pain": "dolor", "lower_pain": "dolor",
    }
    engine = ScientificEvidenceEngine()
    information_list = []
    used = set()
    for column, evidence_name in mapping.items():
        if numeric_values(clinical_data, column).empty:
            continue
        try:
            info = engine.search(evidence_name)
        except Exception:
            info = None
        if not info:
            continue
        label = info.get("label", evidence_name)
        if label in used:
            continue
        used.add(label)
        information_list.append(info)

    if not information_list:
        return

    add_section_heading(document, "6. Contexto científico", level=1)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(
        "La información siguiente contextualiza las variables disponibles y no constituye una valoración diagnóstica."
    )

    for info in information_list:
        add_section_heading(document, info.get("label", "Variable"), level=2)
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(info.get("description", ""))

        rows = [
            ("Significado clínico", info.get("clinical_meaning", "")),
            ("Valores bajos", info.get("low_values", "")),
            ("Valores altos", info.get("high_values", "")),
            ("Recomendación", info.get("recommendation", "")),
            ("Advertencia", info.get("alert", "")),
        ]
        table = document.add_table(rows=0, cols=2)
        for label, value in rows:
            if not value:
                continue
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = str(value)
        for row_index, row in enumerate(table.rows):
            for i, cell in enumerate(row.cells):
                set_cell_shading(cell, LIGHT_BLUE if i == 0 else WHITE)
                set_cell_border(cell)
                set_cell_margins(cell)
                if i == 0 and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(NAVY)

        references = info.get("references", [])
        if references:
            p = document.add_paragraph()
            run = p.add_run("Referencias: ")
            run.bold = True
            p.add_run(" · ".join(str(ref) for ref in references))
            for run in p.runs:
                run.font.size = Pt(7.8)
                run.font.color.rgb = RGBColor.from_string(TEXT_GREY)


def add_temporal_evolution(document: Document, clinical_data: pd.DataFrame) -> None:
    add_section_heading(document, "7. Evolución temporal")
    columns = ["date", *VARIABLES.keys()]
    available = [column for column in columns if column in clinical_data.columns]
    if "date" not in available:
        document.add_paragraph("No hay fechas válidas para construir la evolución temporal.")
        return
    evolution = clinical_data[available].copy()
    evolution["date"] = pd.to_datetime(evolution["date"], errors="coerce", dayfirst=True)
    evolution = evolution[evolution["date"].notna()].sort_values("date").tail(30)
    if evolution.empty:
        document.add_paragraph("No hay registros temporales válidos.")
        return
    evolution["date"] = evolution["date"].dt.strftime("%d/%m/%Y")
    evolution.rename(columns={
        "date": "Fecha", "hr": "FC", "rmssd": "RMSSD", "ln_rmssd": "LnRMSSD",
        "sbp": "PAS", "dbp": "PAD", "spo2": "SpO₂", "sleep": "Sueño",
        "mood": "Ánimo", "stress": "Estrés", "fatigue": "Fatiga",
        "upper_pain": "Dolor sup.", "lower_pain": "Dolor inf.",
    }, inplace=True)

    table = document.add_table(rows=1, cols=len(evolution.columns))
    for i, column in enumerate(evolution.columns):
        table.rows[0].cells[i].text = str(column)
    for _, row_data in evolution.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = format_value(value)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_table(table)


def add_data_quality(document: Document, clinical_data: pd.DataFrame) -> None:
    add_section_heading(document, "8. Calidad y disponibilidad de datos")
    table = document.add_table(rows=1, cols=5)
    headers = ["Variable", "Válidos", "Ausentes", "Disponibilidad", "Valoración"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    for column, (label, _) in VARIABLES.items():
        if column not in clinical_data.columns:
            continue
        valid = int(clinical_data[column].notna().sum())
        missing = int(clinical_data[column].isna().sum())
        total = valid + missing
        availability = valid / total * 100 if total else 0.0
        assessment = "Completa" if availability >= 90 else "Aceptable" if availability >= 70 else "Limitada"
        cells = table.add_row().cells
        outputs = [label, str(valid), str(missing), f"{availability:.1f} %", assessment]
        for i, value in enumerate(outputs):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
    style_table(table)


def trend_interpretation(clinical_data: pd.DataFrame, columns: Iterable[str]) -> str:
    parts: list[str] = []
    for column in columns:
        values = numeric_values(clinical_data, column)
        if values.empty:
            continue
        label, unit = VARIABLES[column]
        first = float(values.iloc[0])
        last = float(values.iloc[-1])
        mean = float(values.mean())
        change = last - first
        status, _, _ = variable_status(column, last)
        direction = "aumentó" if change > 0 else "disminuyó" if change < 0 else "se mantuvo"
        parts.append(
            f"{label}: último valor {last:.1f} {unit}; media {mean:.1f}; "
            f"{direction} {abs(change):.1f} {unit} respecto al primer registro. Estado: {status.lower()}."
        )
    return " ".join(parts) if parts else "No hay datos suficientes para interpretar esta figura."


def add_graphs_section(document: Document, records: pd.DataFrame, clinical_data: pd.DataFrame, participant_folder: Path) -> None:
    add_section_heading(document, "9. Gráficos clínicos")
    figures_folder = participant_folder / "figures"
    try:
        graph_paths = generate_clinical_graphs(records, output_dir=figures_folder)
    except Exception as exc:
        document.add_paragraph(f"No se pudieron generar los gráficos clínicos: {exc}")
        return
    if not graph_paths:
        document.add_paragraph("No existen datos suficientes para generar gráficos.")
        return

    for index, graph_path in enumerate(graph_paths, start=1):
        if index > 1:
            document.add_page_break()
        title = GRAPH_TITLES.get(graph_path.name, graph_path.stem.replace("_", " ").title())
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(title)
        run.bold = True
        run.font.name = "Aptos Display"
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor.from_string(NAVY)

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(graph_path), width=Cm(24.2))

        interpretation = trend_interpretation(clinical_data, GRAPH_VARIABLES.get(graph_path.name, ()))
        table = document.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell, BLUE, "7")
        set_cell_margins(cell, top=110, bottom=110, start=140, end=140)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run("Interpretación del gráfico: ")
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(NAVY)
        p.add_run(interpretation)


def add_conclusions(document: Document, clinical_data: pd.DataFrame) -> None:
    add_section_heading(document, "10. Conclusiones")
    available = [label for column, (label, _) in VARIABLES.items() if not numeric_values(clinical_data, column).empty]
    if not available:
        document.add_paragraph("No existen datos suficientes para elaborar conclusiones.")
        return

    favourable: list[str] = []
    follow_up: list[str] = []
    adverse: list[str] = []
    for column, (label, _) in VARIABLES.items():
        values = numeric_values(clinical_data, column)
        if values.empty:
            continue
        status, _, _ = variable_status(column, float(values.iloc[-1]))
        if status == "FAVORABLE":
            favourable.append(label)
        elif status in {"SEGUIMIENTO", "INFORMATIVO"}:
            follow_up.append(label)
        else:
            adverse.append(label)

    rows = [
        ("Variables analizadas", f"{len(available)}: " + ", ".join(available), NAVY, LIGHT_BLUE),
        ("Situación favorable", ", ".join(favourable) if favourable else "No identificada con los criterios operativos.", GREEN, LIGHT_GREEN),
        ("Seguimiento recomendado", ", ".join(follow_up) if follow_up else "Sin elementos adicionales.", AMBER, LIGHT_AMBER),
        ("Atención prioritaria", ", ".join(adverse) if adverse else "No se identifican valores actuales en categoría desfavorable o de riesgo.", RED, LIGHT_RED),
    ]
    table = document.add_table(rows=0, cols=2)
    for label, text, color, fill in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], color)
        set_cell_shading(cells[1], fill)
        for cell in cells:
            set_cell_border(cell, color, "6")
            set_cell_margins(cell, top=110, bottom=110, start=130, end=130)
        p = cells[0].paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        p = cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(text)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("Consideración final: ")
    run.bold = True
    p.add_run(
        "la interpretación debe considerar la evolución individual, el tratamiento, la medicación, "
        "los síntomas y las condiciones de medición. Los valores aislados no deben utilizarse como diagnóstico."
    )


def normalize_report_mode(report_mode: str) -> str:
    """Normaliza y valida el tipo de informe."""

    mode = (
        str(report_mode or "")
        .strip()
        .upper()
        .replace("-", "_")
    )

    aliases = {
        "1": REPORT_MODE_PRE,
        "PRE": REPORT_MODE_PRE,
        "ANTES": REPORT_MODE_PRE,

        "2": REPORT_MODE_POST,
        "POST": REPORT_MODE_POST,
        "DESPUES": REPORT_MODE_POST,
        "DESPUÉS": REPORT_MODE_POST,

        "3": REPORT_MODE_PRE_POST,
        "PREPOST": REPORT_MODE_PRE_POST,
        "PRE_POST": REPORT_MODE_PRE_POST,
        "EFECTO": REPORT_MODE_PRE_POST,
    }

    normalized = aliases.get(mode, mode)

    if normalized not in REPORT_MODES:
        raise ValueError(
            "Tipo de informe no válido. "
            "Utilice PRE, POST o PRE_POST."
        )

    return normalized


def report_mode_title(report_mode: str) -> str:
    """Título visible del tipo de informe."""

    mode = normalize_report_mode(report_mode)

    return {
        REPORT_MODE_PRE: "ESTADO PRE",
        REPORT_MODE_POST: "ESTADO POST",
        REPORT_MODE_PRE_POST: "EFECTO PRE–POST",
    }[mode]


def report_mode_filename(report_mode: str) -> str:
    """Texto del modo utilizado en el nombre del archivo."""

    mode = normalize_report_mode(report_mode)

    return {
        REPORT_MODE_PRE: "PRE",
        REPORT_MODE_POST: "POST",
        REPORT_MODE_PRE_POST: "PRE_POST",
    }[mode]


def filter_records_for_report(
    records: pd.DataFrame,
    report_mode: str,
) -> pd.DataFrame:
    """Selecciona los registros PRE o POST."""

    mode = normalize_report_mode(report_mode)

    if mode == REPORT_MODE_PRE_POST:
        return records.copy()

    prepared = prepare_form_dataframe(
        records.copy()
    )

    if "moment" not in prepared.columns:
        raise KeyError(
            "No existe la variable 'moment' necesaria "
            "para seleccionar PRE o POST."
        )

    from src.analysis.moment_filter import normalize_moment

    target = (
        1
        if mode == REPORT_MODE_PRE
        else 2
    )

    normalized_moment = prepared[
        "moment"
    ].apply(normalize_moment)

    selected_indexes = prepared.index[
        normalized_moment == target
    ]

    filtered = records.loc[
        records.index.intersection(selected_indexes)
    ].copy()

    if filtered.empty:
        raise ValueError(
            f"No hay registros válidos para generar "
            f"el informe {mode}."
        )

    return filtered


def generate_participant_report(
    records: pd.DataFrame,
    output_dir: Path | str | None = None,
    report_mode: str = REPORT_MODE_PRE,
) -> Path:
    """
    Genera un informe Word PRE, POST o PRE-POST.
    """

    if records is None or records.empty:
        raise ValueError(
            "No hay registros disponibles para generar el informe."
        )

    mode = normalize_report_mode(
        report_mode
    )

    report_records = filter_records_for_report(
        records,
        mode,
    )

    participant = get_current_participant() or {
        "participant_id": "sin_id",
        "name": "Participante",
        "site": "No disponible",
    }

    clinical_data = prepare_form_dataframe(
        report_records.copy()
    )

    clinical_data = clean_clinical_data(
        clinical_data
    )

    if output_dir is None:
        participant_folder = (
            RESULTS_DIR
            / safe_filename(
                participant.get("site")
            )
            / (
                safe_filename(
                    participant.get("participant_id")
                )
                + "_"
                + safe_filename(
                    participant.get("name")
                )
            )
        )
    else:
        participant_folder = Path(
            output_dir
        )

    participant_folder.mkdir(
        parents=True,
        exist_ok=True,
    )

    filename = (
        "Informe_"
        + report_mode_filename(mode)
        + "_"
        + safe_filename(
            participant.get("participant_id")
        )
        + "_"
        + safe_filename(
            participant.get("name")
        )
        + ".docx"
    )

    output_path = (
        participant_folder
        / filename
    )

    document = Document()

    configure_document(
        document
    )

    add_header_footer(
        document
    )

    add_cover(
        document,
        participant,
        report_records,
    )

    mode_paragraph = document.add_paragraph()

    mode_paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    mode_run = mode_paragraph.add_run(
        report_mode_title(mode)
    )

    mode_run.bold = True
    mode_run.font.name = "Aptos Display"
    mode_run.font.size = Pt(18)
    mode_run.font.color.rgb = (
        RGBColor.from_string(BLUE)
    )

    if mode in {
        REPORT_MODE_PRE,
        REPORT_MODE_POST,
    }:
        add_executive_summary(
            document,
            clinical_data,
        )

        add_introduction_and_participant(
            document,
            participant,
            clinical_data,
        )

        add_descriptive_summary(
            document,
            clinical_data,
        )

        add_blood_pressure_section(
            document,
            clinical_data,
        )

        add_wellbeing_section(
            document,
            clinical_data,
        )

        add_scientific_context(
            document,
            clinical_data,
        )

        add_temporal_evolution(
            document,
            clinical_data,
        )

        add_data_quality(
            document,
            clinical_data,
        )

        add_graphs_section(
            document,
            report_records,
            clinical_data,
            participant_folder,
        )

        add_conclusions(
            document,
            clinical_data,
        )

    else:
        add_introduction_and_participant(
            document,
            participant,
            clinical_data,
        )

        add_moment_report_sections(
            document,
            clinical_data,
        )

        add_graphs_section(
            document,
            report_records,
            clinical_data,
            participant_folder,
        )

        add_data_quality(
            document,
            clinical_data,
        )

    document.core_properties.title = (
        "Informe individual - "
        + report_mode_title(mode)
    )

    document.core_properties.subject = (
        PROJECT_NAME
    )

    document.core_properties.author = (
        AUTHOR
    )

    document.core_properties.keywords = (
        "bienestar, cáncer de mama, "
        "PRE, POST, seguimiento longitudinal"
    )

    document.save(
        output_path
    )

    return output_path