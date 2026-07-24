# -*- coding: utf-8 -*-

"""
SportsLabResearch-BreastCancer-Wellbeing-Analyzer

Generador de informe PRE-POST con resumen gráfico clínico.

El script:

1. Solicita una participante.
2. Genera el informe PRE-POST existente.
3. Identifica la última sesión con PRE y POST.
4. Genera una página gráfica de respuesta aguda.
5. Incorpora la página al informe Word.
"""

from __future__ import annotations

import sys
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from main import (
    ID_COLUMN,
    NAME_COLUMN,
    SITE_COLUMN,
    prepare_dataframe,
)
from src.analysis.moment_filter import normalize_moment
from src.cleaning.clinical_cleaning import clean_clinical_data
from src.config.variable_mapping import prepare_form_dataframe
from src.core.session import set_current_participant
from src.reports.report_generator import (
    REPORT_MODE_PRE_POST,
    generate_participant_report,
)


ROOT = Path(__file__).resolve().parent
RESULTS_DIR = ROOT / "results"

NAVY = "#12376B"
BLUE = "#1769AA"
RED = "#D92D20"
GREY = "#7A7A7A"
LIGHT_BLUE = "#EEF5FC"
LIGHT_GREY = "#F5F7FA"
WHITE = "#FFFFFF"
BLACK = "#101828"


VARIABLES = [
    {
        "column": "hr",
        "label": "Frecuencia cardiaca",
        "unit": "bpm",
        "group": "SISTEMA CARDIOVASCULAR",
    },
    {
        "column": "sbp",
        "label": "Presión arterial sistólica",
        "unit": "mmHg",
        "group": "SISTEMA CARDIOVASCULAR",
    },
    {
        "column": "dbp",
        "label": "Presión arterial diastólica",
        "unit": "mmHg",
        "group": "SISTEMA CARDIOVASCULAR",
    },
    {
        "column": "spo2",
        "label": "Saturación de oxígeno",
        "unit": "%",
        "group": "SISTEMA CARDIOVASCULAR",
    },
    {
        "column": "rmssd",
        "label": "RMSSD",
        "unit": "ms",
        "group": "SISTEMA AUTONÓMICO",
    },
    {
        "column": "ln_rmssd",
        "label": "LnRMSSD",
        "unit": "ln(ms)",
        "group": "SISTEMA AUTONÓMICO",
    },
    {
        "column": "sleep",
        "label": "Calidad del sueño",
        "unit": "puntos",
        "group": "BIENESTAR Y BIORRETROALIMENTACIÓN",
    },
    {
        "column": "mood",
        "label": "Estado de ánimo",
        "unit": "puntos",
        "group": "BIENESTAR Y BIORRETROALIMENTACIÓN",
    },
    {
        "column": "stress",
        "label": "Estrés",
        "unit": "puntos",
        "group": "BIENESTAR Y BIORRETROALIMENTACIÓN",
    },
    {
        "column": "fatigue",
        "label": "Fatiga",
        "unit": "puntos",
        "group": "BIENESTAR Y BIORRETROALIMENTACIÓN",
    },
    {
        "column": "upper_pain",
        "label": "Dolor cuerpo superior",
        "unit": "puntos",
        "group": "DOLOR",
    },
    {
        "column": "lower_pain",
        "label": "Dolor cuerpo inferior",
        "unit": "puntos",
        "group": "DOLOR",
    },
]


INTERPRETATIONS = {
    "hr": {
        "increase": (
            "Aumento de la frecuencia cardiaca tras la sesión. "
            "Interpretar según intensidad y tiempo de recuperación."
        ),
        "decrease": (
            "Disminución de la frecuencia cardiaca tras la sesión. "
            "Interpretar según el momento exacto de la medición."
        ),
        "same": "La frecuencia cardiaca no mostró cambios relevantes.",
    },
    "sbp": {
        "increase": (
            "Aumento de la presión arterial sistólica tras la sesión. "
            "Interpretar junto con la carga y las condiciones de medición."
        ),
        "decrease": (
            "Disminución de la presión arterial sistólica tras la sesión. "
            "Puede reflejar recuperación o respuesta postejercicio."
        ),
        "same": "La presión arterial sistólica se mantuvo estable.",
    },
    "dbp": {
        "increase": (
            "Aumento de la presión arterial diastólica tras la sesión. "
            "Interpretar dentro del contexto hemodinámico individual."
        ),
        "decrease": (
            "Disminución de la presión arterial diastólica tras la sesión. "
            "Interpretar según síntomas y condiciones de registro."
        ),
        "same": "La presión arterial diastólica se mantuvo estable.",
    },
    "spo2": {
        "increase": "Aumento de la saturación de oxígeno tras la sesión.",
        "decrease": (
            "Disminución de la saturación de oxígeno tras la sesión. "
            "Valorar magnitud, persistencia y calidad de la medición."
        ),
        "same": "La saturación de oxígeno se mantuvo estable.",
    },
    "rmssd": {
        "increase": (
            "Aumento de RMSSD tras la sesión. "
            "Interpretar según recuperación y condiciones del registro."
        ),
        "decrease": (
            "Descenso de RMSSD tras la sesión. "
            "Respuesta compatible con activación autonómica aguda."
        ),
        "same": "RMSSD no mostró cambios relevantes.",
    },
    "ln_rmssd": {
        "increase": (
            "Aumento de LnRMSSD tras la sesión. "
            "Interpretar junto con RMSSD y la calidad de los intervalos RR."
        ),
        "decrease": (
            "Descenso de LnRMSSD tras la sesión. "
            "Compatible con una respuesta autonómica aguda."
        ),
        "same": "LnRMSSD no mostró cambios relevantes.",
    },
    "sleep": {
        "increase": "Aumento de la puntuación de sueño registrada.",
        "decrease": "Disminución de la puntuación de sueño registrada.",
        "same": "La puntuación de sueño se mantuvo estable.",
    },
    "mood": {
        "increase": "Aumento del estado de ánimo percibido.",
        "decrease": "Disminución del estado de ánimo percibido.",
        "same": "El estado de ánimo percibido se mantuvo estable.",
    },
    "stress": {
        "increase": "Incremento del estrés percibido tras la sesión.",
        "decrease": "Disminución del estrés percibido tras la sesión.",
        "same": "El estrés percibido se mantuvo estable.",
    },
    "fatigue": {
        "increase": (
            "Aumento de la fatiga percibida tras la sesión. "
            "Interpretar en relación con la carga realizada."
        ),
        "decrease": "Disminución de la fatiga percibida tras la sesión.",
        "same": "La fatiga percibida se mantuvo estable.",
    },
    "upper_pain": {
        "increase": (
            "Aumento del dolor en el cuerpo superior tras la sesión. "
            "Valorar magnitud, persistencia e impacto funcional."
        ),
        "decrease": "Disminución del dolor en el cuerpo superior.",
        "same": "Sin cambios en el dolor del cuerpo superior.",
    },
    "lower_pain": {
        "increase": (
            "Aumento del dolor en el cuerpo inferior tras la sesión. "
            "Valorar magnitud, persistencia e impacto funcional."
        ),
        "decrease": "Disminución del dolor en el cuerpo inferior.",
        "same": "Sin cambios en el dolor del cuerpo inferior.",
    },
}


def safe_filename(value: Any) -> str:
    text = str(value or "").strip()

    for old, new in {
        " ": "_",
        "/": "-",
        "\\": "-",
        ":": "-",
        "*": "",
        "?": "",
        '"': "",
        "<": "",
        ">": "",
        "|": "",
    }.items():
        text = text.replace(old, new)

    return text or "participante"


def most_frequent_value(
    dataframe: pd.DataFrame,
    column: str,
    default: str,
) -> str:
    if column not in dataframe.columns:
        return default

    values = (
        dataframe[column]
        .dropna()
        .astype(str)
        .str.strip()
    )

    values = values[values != ""]

    if values.empty:
        return default

    return str(values.value_counts().index[0])


def numeric_value(
    row: pd.Series,
    column: str,
) -> float:
    if column not in row.index:
        return np.nan

    return pd.to_numeric(
        pd.Series([row[column]]),
        errors="coerce",
    ).iloc[0]


def format_number(
    value: float,
    decimals: int = 1,
) -> str:
    if pd.isna(value):
        return "-"

    if float(value).is_integer():
        return str(int(value))

    return f"{float(value):.{decimals}f}"


def format_datetime(value: Any) -> str:
    date = pd.to_datetime(
        value,
        errors="coerce",
        dayfirst=True,
    )

    if pd.isna(date):
        return "No disponible"

    if date.hour == 0 and date.minute == 0:
        return date.strftime("%d/%m/%Y")

    return date.strftime("%d/%m/%Y %H:%M")


def find_datetime_column(
    dataframe: pd.DataFrame,
) -> str:
    candidates = [
        "datetime",
        "timestamp",
        "date_time",
        "fecha_hora",
        "recorded_at",
        "date",
    ]

    for column in candidates:
        if column in dataframe.columns:
            return column

    raise KeyError(
        "No se encontró una columna de fecha válida."
    )


def select_latest_pre_post_pair(
    clinical_data: pd.DataFrame,
) -> tuple[pd.Series, pd.Series, str]:
    if "moment" not in clinical_data.columns:
        raise KeyError(
            "No existe la columna 'moment'."
        )

    dataframe = clinical_data.copy()

    date_column = find_datetime_column(
        dataframe
    )

    dataframe["_datetime_pair"] = pd.to_datetime(
        dataframe[date_column],
        errors="coerce",
        dayfirst=True,
    )

    dataframe["_moment_normalized"] = dataframe[
        "moment"
    ].apply(normalize_moment)

    dataframe = dataframe[
        dataframe["_moment_normalized"].isin([1, 2])
        & dataframe["_datetime_pair"].notna()
    ].copy()

    if dataframe.empty:
        raise ValueError(
            "No existen registros PRE o POST válidos."
        )

    dataframe["_session_date"] = (
        dataframe["_datetime_pair"].dt.date
    )

    valid_dates = []

    for session_date, group in dataframe.groupby(
        "_session_date"
    ):
        moments = set(
            group["_moment_normalized"].dropna()
        )

        if 1 in moments and 2 in moments:
            valid_dates.append(session_date)

    if not valid_dates:
        raise ValueError(
            "No existe ninguna fecha con registros PRE y POST."
        )

    latest_date = max(valid_dates)

    session = dataframe[
        dataframe["_session_date"] == latest_date
    ].copy()

    pre_rows = session[
        session["_moment_normalized"] == 1
    ].sort_values("_datetime_pair")

    post_rows = session[
        session["_moment_normalized"] == 2
    ].sort_values("_datetime_pair")

    if pre_rows.empty or post_rows.empty:
        raise ValueError(
            "No se pudo construir el último par PRE-POST."
        )

    pre_row = pre_rows.iloc[-1]
    post_row = post_rows.iloc[-1]

    return pre_row, post_row, date_column


def change_interpretation(
    column: str,
    delta: float,
) -> str:
    if pd.isna(delta):
        return "No evaluable por ausencia de datos."

    tolerance = 1e-9

    if delta > tolerance:
        direction = "increase"
    elif delta < -tolerance:
        direction = "decrease"
    else:
        direction = "same"

    configuration = INTERPRETATIONS.get(
        column,
        {},
    )

    return configuration.get(
        direction,
        "Cambio observado tras la sesión.",
    )


def collect_graph_rows(
    pre_row: pd.Series,
    post_row: pd.Series,
) -> list[dict[str, Any]]:
    rows = []

    for variable in VARIABLES:
        column = variable["column"]

        pre_value = numeric_value(
            pre_row,
            column,
        )

        post_value = numeric_value(
            post_row,
            column,
        )

        if pd.isna(pre_value) and pd.isna(post_value):
            continue

        delta = (
            post_value - pre_value
            if pd.notna(pre_value) and pd.notna(post_value)
            else np.nan
        )

        delta_percent = (
            delta / pre_value * 100
            if pd.notna(delta)
            and pd.notna(pre_value)
            and pre_value != 0
            else np.nan
        )

        rows.append({
            **variable,
            "pre": pre_value,
            "post": post_value,
            "delta": delta,
            "delta_percent": delta_percent,
            "interpretation": change_interpretation(
                column,
                delta,
            ),
        })

    return rows


def draw_text(
    axis,
    x: float,
    y: float,
    text: str,
    size: float = 9,
    weight: str = "normal",
    color: str = BLACK,
    horizontal: str = "left",
    vertical: str = "center",
) -> None:
    axis.text(
        x,
        y,
        text,
        fontsize=size,
        fontweight=weight,
        color=color,
        ha=horizontal,
        va=vertical,
        family="DejaVu Sans",
        wrap=True,
    )


def generate_visual_summary(
    participant_id: str,
    participant_name: str,
    participant_site: str,
    pre_row: pd.Series,
    post_row: pd.Series,
    date_column: str,
    output_path: Path,
) -> Path:
    rows = collect_graph_rows(
        pre_row,
        post_row,
    )

    if not rows:
        raise ValueError(
            "No existen variables válidas para el gráfico PRE-POST."
        )

    pre_date = format_datetime(
        pre_row[date_column]
    )

    post_date = format_datetime(
        post_row[date_column]
    )

    session_date = pd.to_datetime(
        pre_row[date_column],
        errors="coerce",
        dayfirst=True,
    )

    session_text = (
        session_date.strftime("%d/%m/%Y")
        if pd.notna(session_date)
        else "No disponible"
    )

    figure = plt.figure(
        figsize=(11.7, 16.5),
        dpi=170,
        facecolor=WHITE,
    )

    axis = figure.add_axes(
        [0, 0, 1, 1]
    )

    axis.set_xlim(0, 1)
    axis.set_ylim(0, 1)
    axis.axis("off")

    draw_text(
        axis,
        0.5,
        0.974,
        "RESPUESTA AGUDA A LA SESIÓN (PRE–POST)",
        size=19,
        weight="bold",
        color=NAVY,
        horizontal="center",
    )

    axis.add_patch(
        plt.Rectangle(
            (0.025, 0.885),
            0.95,
            0.065,
            facecolor=WHITE,
            edgecolor="#A9C4E4",
            linewidth=1.2,
        )
    )

    draw_text(
        axis,
        0.05,
        0.93,
        f"Participante:  {participant_id} – {participant_name}",
        size=10,
        weight="bold",
        color=NAVY,
    )

    draw_text(
        axis,
        0.05,
        0.907,
        f"Sede:  {participant_site}",
        size=9,
    )

    draw_text(
        axis,
        0.05,
        0.888,
        f"Fecha de la sesión:  {session_text}",
        size=9,
    )

    draw_text(
        axis,
        0.63,
        0.93,
        f"PRE:   {pre_date}",
        size=9.5,
        weight="bold",
        color=NAVY,
    )

    draw_text(
        axis,
        0.63,
        0.907,
        f"POST:  {post_date}",
        size=9.5,
        weight="bold",
        color=NAVY,
    )

    header_y = 0.855
    header_height = 0.04

    axis.add_patch(
        plt.Rectangle(
            (0.025, header_y),
            0.95,
            header_height,
            facecolor=NAVY,
            edgecolor=NAVY,
        )
    )

    draw_text(
        axis,
        0.055,
        header_y + header_height / 2,
        "VARIABLE",
        size=9,
        weight="bold",
        color=WHITE,
    )

    draw_text(
        axis,
        0.375,
        header_y + header_height / 2,
        "PRE                         POST",
        size=9,
        weight="bold",
        color=WHITE,
        horizontal="center",
    )

    draw_text(
        axis,
        0.675,
        header_y + header_height / 2,
        "CAMBIO",
        size=9,
        weight="bold",
        color=WHITE,
        horizontal="center",
    )

    draw_text(
        axis,
        0.84,
        header_y + header_height / 2,
        "INTERPRETACIÓN",
        size=9,
        weight="bold",
        color=WHITE,
        horizontal="center",
    )

    groups = []

    for row in rows:
        if row["group"] not in groups:
            groups.append(row["group"])

    available_height = 0.735
    group_header_height = 0.027
    row_height = (
        available_height
        - len(groups) * group_header_height
    ) / max(len(rows), 1)

    current_y = header_y

    for group in groups:
        group_rows = [
            row
            for row in rows
            if row["group"] == group
        ]

        current_y -= group_header_height

        axis.add_patch(
            plt.Rectangle(
                (0.025, current_y),
                0.95,
                group_header_height,
                facecolor=LIGHT_BLUE,
                edgecolor="#C7D7EA",
                linewidth=0.7,
            )
        )

        draw_text(
            axis,
            0.05,
            current_y + group_header_height / 2,
            group,
            size=8.5,
            weight="bold",
            color=NAVY,
        )

        for row_index, row in enumerate(group_rows):
            current_y -= row_height

            fill = (
                WHITE
                if row_index % 2 == 0
                else LIGHT_GREY
            )

            axis.add_patch(
                plt.Rectangle(
                    (0.025, current_y),
                    0.95,
                    row_height,
                    facecolor=fill,
                    edgecolor="#D9E0E8",
                    linewidth=0.6,
                )
            )

            label = (
                f"{row['label']} ({row['unit']})"
            )

            draw_text(
                axis,
                0.04,
                current_y + row_height / 2,
                label,
                size=7.8,
                weight="bold",
            )

            pre_text = format_number(
                row["pre"]
            )

            post_text = format_number(
                row["post"]
            )

            y_center = (
                current_y + row_height / 2
            )

            x_pre = 0.30
            x_post = 0.54

            draw_text(
                axis,
                x_pre - 0.024,
                y_center,
                pre_text,
                size=8.5,
                weight="bold",
                horizontal="right",
            )

            draw_text(
                axis,
                x_post + 0.024,
                y_center,
                post_text,
                size=8.5,
                weight="bold",
                horizontal="left",
            )

            axis.plot(
                [x_pre, x_post],
                [y_center, y_center],
                color=BLUE,
                linewidth=1.7,
                solid_capstyle="round",
            )

            axis.scatter(
                [x_pre],
                [y_center],
                s=28,
                color=BLUE,
                zorder=3,
            )

            if pd.isna(row["delta"]):
                post_color = GREY
            elif row["delta"] > 0:
                post_color = RED
            elif row["delta"] < 0:
                post_color = BLUE
            else:
                post_color = GREY

            axis.scatter(
                [x_post],
                [y_center],
                s=28,
                color=post_color,
                zorder=3,
            )

            delta_text = (
                "-"
                if pd.isna(row["delta"])
                else f"{row['delta']:+.1f}"
            )

            percent_text = (
                "-"
                if pd.isna(row["delta_percent"])
                else f"{row['delta_percent']:+.1f} %"
            )

            draw_text(
                axis,
                0.66,
                y_center + row_height * 0.15,
                delta_text,
                size=8,
                weight="bold",
                color=post_color,
                horizontal="center",
            )

            draw_text(
                axis,
                0.71,
                y_center - row_height * 0.15,
                percent_text,
                size=7.4,
                color=post_color,
                horizontal="center",
            )

            draw_text(
                axis,
                0.755,
                y_center,
                row["interpretation"],
                size=6.7,
                color=BLACK,
            )

    note_y = 0.035

    axis.add_patch(
        plt.Rectangle(
            (0.025, note_y),
            0.95,
            0.065,
            facecolor=LIGHT_BLUE,
            edgecolor="#A9C4E4",
            linewidth=1.0,
        )
    )

    draw_text(
        axis,
        0.05,
        note_y + 0.047,
        "REFERENCIA DE CAMBIOS",
        size=8.2,
        weight="bold",
        color=NAVY,
    )

    draw_text(
        axis,
        0.05,
        note_y + 0.028,
        "Rojo: aumento · Azul: disminución · Gris: sin cambio",
        size=7.5,
    )

    draw_text(
        axis,
        0.05,
        note_y + 0.012,
        "Δ %: cambio porcentual respecto al valor PRE",
        size=7.5,
    )

    draw_text(
        axis,
        0.48,
        note_y + 0.047,
        "NOTA CLÍNICA",
        size=8.2,
        weight="bold",
        color=NAVY,
    )

    draw_text(
        axis,
        0.48,
        note_y + 0.023,
        (
            "Los cambios representan la respuesta aguda de la última "
            "sesión con un par PRE–POST válido. La interpretación debe "
            "considerar la carga, el tratamiento, los síntomas y el "
            "momento exacto de registro."
        ),
        size=7.2,
    )

    draw_text(
        axis,
        0.5,
        0.012,
        (
            "Informe generado automáticamente por SportsLabResearch – "
            "Breast Cancer Wellbeing Analyzer"
        ),
        size=7.2,
        color=NAVY,
        horizontal="center",
    )

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    figure.savefig(
        output_path,
        dpi=170,
        facecolor=WHITE,
        bbox_inches="tight",
        pad_inches=0.08,
    )

    plt.close(figure)

    return output_path


def append_visual_to_report(
    source_report: Path,
    image_path: Path,
    output_report: Path,
) -> Path:
    document = Document(
        source_report
    )

    section = document.add_section(
        WD_SECTION.NEW_PAGE
    )

    section.orientation = (
        WD_ORIENT.PORTRAIT
    )

    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.6)
    section.bottom_margin = Cm(0.6)
    section.left_margin = Cm(0.7)
    section.right_margin = Cm(0.7)

    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0

    paragraph.add_run().add_picture(
        str(image_path),
        width=Cm(19.3),
    )

    document.save(
        output_report
    )

    return output_report


def main() -> int:
    print("=" * 92)
    print("INFORME PRE-POST CON RESUMEN GRÁFICO")
    print("=" * 92)

    participant_id = input(
        "\nIntroduzca el identificador de la participante: "
    ).strip()

    if not participant_id:
        print(
            "\nNo se ha introducido ningún identificador."
        )
        return 1

    try:
        dataframe = prepare_dataframe()
    except Exception as exc:
        print("\n[ERROR AL CARGAR LOS DATOS]")
        print(str(exc))
        return 1

    records = dataframe[
        dataframe[ID_COLUMN]
        .astype(str)
        .str.strip()
        == participant_id
    ].copy()

    if records.empty:
        print(
            f"\nNo se encontraron registros para "
            f"el identificador {participant_id}."
        )
        return 1

    participant_name = most_frequent_value(
        records,
        NAME_COLUMN,
        "Participante",
    )

    participant_site = most_frequent_value(
        records,
        SITE_COLUMN,
        "No disponible",
    )

    set_current_participant(
        participant_id=participant_id,
        name=participant_name,
        site=participant_site,
        dataframe=records,
    )

    print()
    print("PARTICIPANTE")
    print("-" * 92)
    print(f"Identificador : {participant_id}")
    print(f"Nombre        : {participant_name}")
    print(f"Sede          : {participant_site}")
    print(f"Registros     : {len(records)}")

    try:
        print(
            "\nGenerando informe PRE-POST..."
        )

        source_report = generate_participant_report(
            records=records,
            report_mode=REPORT_MODE_PRE_POST,
        )

        clinical_data = prepare_form_dataframe(
            records.copy()
        )

        clinical_data = clean_clinical_data(
            clinical_data
        )

        pre_row, post_row, date_column = (
            select_latest_pre_post_pair(
                clinical_data
            )
        )

        participant_folder = (
            Path(source_report).parent
        )

        figures_folder = (
            participant_folder
            / "figures"
        )

        image_path = (
            figures_folder
            / (
                "Resumen_grafico_PRE_POST_"
                + safe_filename(participant_id)
                + ".png"
            )
        )

        generate_visual_summary(
            participant_id=participant_id,
            participant_name=participant_name,
            participant_site=participant_site,
            pre_row=pre_row,
            post_row=post_row,
            date_column=date_column,
            output_path=image_path,
        )

        output_report = (
            participant_folder
            / (
                "Informe_PRE_POST_GRAFICO_"
                + safe_filename(participant_id)
                + "_"
                + safe_filename(participant_name)
                + ".docx"
            )
        )

        append_visual_to_report(
            source_report=Path(source_report),
            image_path=image_path,
            output_report=output_report,
        )

    except Exception as exc:
        print(
            "\n[ERROR AL GENERAR EL INFORME]"
        )
        print(str(exc))
        return 1

    print()
    print("=" * 92)
    print("INFORME GENERADO CORRECTAMENTE")
    print("=" * 92)
    print(f"Informe : {output_report}")
    print(f"Gráfico : {image_path}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
